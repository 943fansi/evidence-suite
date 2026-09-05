#!/usr/bin/env python3
r"""Export evidence-driven draft Markdown to a properly typeset .docx.

Why: real deliverables in this pipeline (专利申请草案/交底书、学位论文、报告)
must ship as Word documents with Chinese-document typography — line spacing,
2-char first-line indent, 宋体/黑体 font pairing, page-break-safe tables. Plain
pandoc/markdown conversion produces none of that.

Typography defaults (Chinese legal/thesis convention, overridable by flags):
  - Page: A4, margins 上下 2.54cm / 左右 3.17cm
  - Title (# first line): 黑体 三号(16pt), centered
  - H2: 黑体 四号(14pt)  | H3: 黑体 小四(12pt), left
  - Body: 宋体 小四(12pt), western font Times New Roman, 1.5x line spacing,
    first-line indent exactly 2 characters (w:firstLineChars=200 — Word-native
    char units, survives font-size changes), justified
  - Numbered patent claims ("1. ..." etc.) and （注：…）lines: same body style
  - Blockquote (推理链 etc.): 楷体 五号(10.5pt), left indent, no first-line indent
  - Mermaid code blocks: rendered to PNG (local mmdc first, mermaid.ink /img/ fallback)
    and embedded as centered figures (same local-first policy as export_pdf.py);
    on render failure a visible placeholder note is emitted instead
  - Markdown tables → Word tables: 宋体 五号(10.5pt), grid borders,
    header bold, rows kept intact across pages (cantSplit)

Usage:
  python scripts/export_docx.py 11_定稿.md                     # → 11_定稿.docx
  python scripts/export_docx.py 11_定稿.md -o 草案.docx --line-spacing 1.5
  python scripts/export_docx.py thesis.md --body-size 12 --no-indent
  python scripts/export_docx.py 11_定稿.md --mermaid-engine local   # forbid remote render

Dependencies: python-docx (pip install python-docx) — optional, lazy import.
Exit codes: 0 ok; 1 missing python-docx / conversion failure; 2 usage error.
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

from mermaid_render import render_mermaid

from evidence_boundary import BOUNDARY_NOTICE, review_kind_label

MD_TABLE_SEP = re.compile(r"^\s*\|?[\s:|-]+\|?\s*$")


def _ensure_utf8_streams() -> None:
    for stream in (sys.stdout, sys.stderr):
        if hasattr(stream, "reconfigure"):
            try:
                stream.reconfigure(encoding="utf-8", errors="replace")
            except Exception:
                pass


def _set_eastasia(run, eastasia: str) -> None:
    """python-docx only sets w:ascii/w:hAnsi via font.name; CJK needs w:eastAsia."""
    from docx.oxml.ns import qn
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), eastasia)


def _set_first_line_chars(paragraph, chars: int = 200) -> None:
    """2-char first-line indent in Word-native character units."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = paragraph._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        jc = pPr.find(qn("w:jc"))
        if jc is not None:
            jc.addprevious(ind)  # CT_PPr 序：spacing → ind → jc，插到 jc 前
        else:
            pPr.append(ind)
    ind.set(qn("w:firstLineChars"), str(chars))
    ind.set(qn("w:firstLine"), "0")  # firstLineChars takes precedence in Word


def _add_runs(paragraph, text: str, *, eastasia: str, western: str,
              size_pt: float, bold_all: bool = False) -> None:
    """Add runs with **bold** and *italic* inline markdown support."""
    for token in re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text):
        if not token:
            continue
        bold = token.startswith("**") and token.endswith("**") and len(token) > 4
        italic = (not bold) and token.startswith("*") and token.endswith("*") and len(token) > 2
        content = token[2:-2] if bold else (token[1:-1] if italic else token)
        content = content.replace("`", "")
        run = paragraph.add_run(content)
        run.font.name = western
        _set_eastasia(run, eastasia)
        run.font.size = None if size_pt is None else _pt(size_pt)
        run.bold = bold or bold_all
        run.italic = italic


def _pt(v: float):
    from docx.shared import Pt
    return Pt(v)


def convert(md_path: Path, out_path: Path, *, body_size: float, line_spacing: float,
            indent: bool, heiti: str, songti: str, kaiti: str, western: str,
            mermaid_engine: str = "auto", manifest_path: Path | None = None) -> int:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.shared import Cm
    except ImportError:
        print("ERROR: python-docx not installed — pip install python-docx", file=sys.stderr)
        return 1

    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    # Page setup: A4 + Chinese-standard margins
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.top_margin = sec.bottom_margin = Cm(2.54)
    sec.left_margin = sec.right_margin = Cm(3.17)

    # Normal style baseline
    normal = doc.styles["Normal"]
    normal.font.name = western
    normal.font.size = _pt(body_size)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), songti)
    normal.paragraph_format.line_spacing = line_spacing
    normal.paragraph_format.space_before = _pt(0)
    normal.paragraph_format.space_after = _pt(0)

    def add_para(text: str, *, eastasia: str = songti, size: float = body_size,
                 bold: bool = False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                 first_indent: bool = False, left_indent_cm: float | None = None):
        p = doc.add_paragraph()
        p.alignment = align
        pf = p.paragraph_format
        pf.line_spacing = line_spacing
        if first_indent and indent:
            _set_first_line_chars(p, 200)
        if left_indent_cm is not None:
            pf.left_indent = Cm(left_indent_cm)
        if text:
            _add_runs(p, text, eastasia=eastasia, western=western,
                      size_pt=size, bold_all=bold)
        return p

    def add_table(rows: list[list[str]]):
        cols = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        for i, row in enumerate(rows):
            # keep each row intact across page breaks (cantSplit)
            trPr = table.rows[i]._tr.get_or_add_trPr()
            cant = OxmlElement("w:cantSplit")
            trPr.append(cant)
            for j in range(cols):
                cell = table.cell(i, j)
                text = row[j] if j < len(row) else ""
                cp = cell.paragraphs[0]
                cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                cp.paragraph_format.line_spacing = line_spacing
                _add_runs(cp, text, eastasia=songti, western=western,
                          size_pt=body_size - 1.5, bold_all=(i == 0))
        doc.add_paragraph()  # spacing after table

    i, n = 0, len(lines)
    title_done = False
    mermaid_counter = [0]
    figures_dir = md_path.parent / "figures"
    while i < n:
        raw = lines[i]
        line = raw.rstrip()

        if not line.strip():
            i += 1
            continue

        # fenced code blocks → mermaid PNG embedding, else placeholder note
        if line.strip().startswith("```"):
            fence = line.strip()[3:].strip()
            block: list[str] = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1
            if fence.lower() == "mermaid":
                code = "\n".join(block).strip()
                if code:
                    fig_seq = mermaid_counter[0] + 1
                    img_path = figures_dir / f"mermaid_docx_{fig_seq}.png"
                    figures_dir.mkdir(parents=True, exist_ok=True)
                    data = render_mermaid(code, img_path, engine=mermaid_engine, fmt="png")
                    if data:
                        mermaid_counter[0] = fig_seq
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.line_spacing = line_spacing
                        p.add_run().add_picture(str(img_path), width=Cm(14.5))
                        continue
                    add_para(f"〔mermaid 图 {fig_seq} 渲染失败——请精简节点标签（≤~12 个中文字符）、"
                             f"或安装 mermaid-cli 后重跑；原始代码见 export_pdf.py 渲染路径〕",
                             eastasia=kaiti, size=body_size - 1.5,
                             align=WD_ALIGN_PARAGRAPH.LEFT, left_indent_cm=0.74)
                    continue
            add_para(f"〔{fence or 'code'} 块未嵌入——图表请由 export_pdf.py 渲染或另附〕",
                     eastasia=kaiti, size=body_size - 1.5,
                     align=WD_ALIGN_PARAGRAPH.LEFT, left_indent_cm=0.74)
            continue

        # markdown tables
        if line.lstrip().startswith("|") and i + 1 < n and MD_TABLE_SEP.match(lines[i + 1]):
            rows = []
            header = [c.strip() for c in line.strip().strip("|").split("|")]
            rows.append(header)
            i += 2
            while i < n and lines[i].lstrip().startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            add_table(rows)
            continue

        # headings
        m = re.match(r"^(#{1,4})\s+(.*)$", line)
        if m:
            level, text = len(m.group(1)), m.group(2).strip()
            if level == 1 and not title_done:
                # document title: 黑体三号居中
                add_para(text, eastasia=heiti, size=16, bold=True,
                         align=WD_ALIGN_PARAGRAPH.CENTER)
                title_done = True
            elif level == 1:
                add_para(text, eastasia=heiti, size=14, bold=True,
                         align=WD_ALIGN_PARAGRAPH.CENTER)
            elif level == 2:
                pf = add_para(text, eastasia=heiti, size=14, bold=True,
                              align=WD_ALIGN_PARAGRAPH.LEFT).paragraph_format
                pf.space_before = _pt(12)
                pf.space_after = _pt(6)
            else:
                pf = add_para(text, eastasia=heiti, size=body_size, bold=True,
                              align=WD_ALIGN_PARAGRAPH.LEFT).paragraph_format
                pf.space_before = _pt(6)
            i += 1
            continue

        # blockquote → 楷体 annotation block
        if line.lstrip().startswith(">"):
            quote = [re.sub(r"^\s*>\s?", "", l) for l in [line]]
            i += 1
            while i < n and lines[i].lstrip().startswith(">"):
                quote.append(re.sub(r"^\s*>\s?", "", lines[i]))
                i += 1
            text = " ".join(q for q in quote if q.strip())
            add_para(text, eastasia=kaiti, size=body_size - 1.5,
                     align=WD_ALIGN_PARAGRAPH.LEFT, left_indent_cm=0.74)
            continue

        # list items
        m = re.match(r"^\s*[-*]\s+(.*)$", line)
        if m:
            add_para("• " + m.group(1), align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     first_indent=False, left_indent_cm=0.74)
            i += 1
            continue

        # normal paragraph (claims "N. …", （注：…）, body prose) → 2-char indent
        add_para(line.strip(), first_indent=True)
        i += 1

    # review-kind footer + capability boundary disclosure (PR-01)
    review_kind = None
    if manifest_path is not None:
        try:
            mdata = json.loads(manifest_path.read_text(encoding="utf-8"))
            review_kind = mdata.get("review_kind")
        except Exception:
            review_kind = None
    label = review_kind_label(review_kind) if review_kind else "未声明（--manifest 提供时自动标注）"
    footer = add_para(f"评审类型：{label}　{BOUNDARY_NOTICE}",
                      eastasia=kaiti, size=body_size - 3,
                      align=WD_ALIGN_PARAGRAPH.LEFT)
    footer.paragraph_format.space_before = _pt(18)

    doc.save(str(out_path))
    print(f"Done: {out_path} ({out_path.stat().st_size:,} bytes, "
          f"body {body_size}pt/{line_spacing}x, indent {'2字符' if indent else 'off'})")
    return 0


def main() -> int:
    _ensure_utf8_streams()
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input", type=Path, help="Draft markdown (e.g., 11_定稿.md)")
    parser.add_argument("-o", "--output", type=Path, default=None,
                        help="Output .docx (default: input stem + .docx)")
    parser.add_argument("--body-size", type=float, default=12.0,
                        help="Body font size in pt (default 12 = 小四)")
    parser.add_argument("--line-spacing", type=float, default=1.5,
                        help="Line spacing multiple (default 1.5)")
    parser.add_argument("--no-indent", action="store_true",
                        help="Disable 2-char first-line indent")
    parser.add_argument("--heiti", default="黑体", help="Heading CJK font (default 黑体)")
    parser.add_argument("--songti", default="宋体", help="Body CJK font (default 宋体)")
    parser.add_argument("--kaiti", default="楷体", help="Annotation CJK font (default 楷体)")
    parser.add_argument("--western", default="Times New Roman", help="Western font")
    parser.add_argument("--mermaid-engine", choices=["auto", "local", "remote"], default="auto",
                        help="Mermaid renderer: auto (local mmdc first, mermaid.ink /img/ "
                             "fallback), local (mmdc only, no network), remote (mermaid.ink only)")
    parser.add_argument("--manifest", type=Path, default=None,
                        help="evidence_manifest.json → 文末标注 review_kind（评审类型）"
                             "与能力边界声明")
    args = parser.parse_args()

    if not args.input.exists():
        print(f"File not found: {args.input}", file=sys.stderr)
        return 2
    out = args.output or args.input.with_suffix(".docx")
    if out.suffix.lower() != ".docx":
        out = out.with_suffix(".docx")
    return convert(args.input, out, body_size=args.body_size,
                   line_spacing=args.line_spacing, indent=not args.no_indent,
                   heiti=args.heiti, songti=args.songti, kaiti=args.kaiti,
                   western=args.western, mermaid_engine=args.mermaid_engine,
                   manifest_path=args.manifest)


if __name__ == "__main__":
    raise SystemExit(main())
